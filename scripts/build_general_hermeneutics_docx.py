from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "general-hermeneutics-no-james-2.md"
OUTPUT = ROOT / "general-hermeneutics-no-james-2.docx"

NAVY = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x5D, 0x6B, 0x7A)
LIGHT_BLUE = "E8EEF5"
PAPER = "F7F4ED"

SECTIONS = [
    "General Hermeneutics Overview",
    "Preparation",
    "Observation",
    "Interpretation",
    "Imagination",
    "Application",
]

MINOR_LABELS = {
    "WHAT IT IS",
    "WHY IT MATTERS",
    "HOW TO BEGIN",
    "HOW TO DO IT",
    "HOW TO ENTER THE SCENE",
    "DETECTIVE",
    "BIBLE STUDY",
    "PAUSE",
    "NOTICE",
    "TEST",
    "REBUILD",
    "ACT",
    "ASK",
    "LOOK FOR",
    "AVOID",
    "AUTHOR",
    "ORIGINAL AUDIENCE",
    "PRIMARY THEME",
    "SPECIFIC OCCASION & PURPOSE",
    "GUARDRAIL",
    "PAUSE AND PRAY",
}

MAJOR_LABELS = {
    "HOW THE INVESTIGATION MOVES",
    "GENRE GUIDE PAGES",
    "LOCAL-CONTEXT WORKFLOW: START NEAR",
    "CONTEXT LADDER: HOW TO USE IT",
    "QUESTION CHECK",
    "HOW TO DO A WORD STUDY",
    "HELPFUL RESOURCES",
    "EXEGETICAL FALLACIES",
    "APPLICATION GUARDRAILS",
    "S.P.A.C.E.P.E.T.S.",
    "CIRCLES OF APPLICATION",
}

TITLE_HEADINGS = {
    "The Detective Method",
    "What Is Hermeneutics?",
    "Who Controls the Meaning?",
    "Two Ways to Come to a Verse",
    "The Inductive Bible Study Method",
    "Start Here",
    "Bible Translations (Formal vs. Functional)",
    "Read and Reread Before Outlining",
    "Passage Boundaries, Thought Units, and Outlines",
    "Make Your Observations: The Complete Catalog. Use the full catalog below as your training grid. Do not shorten the search too early; let the obvious details become the floor for everything subtler.",
    "66-Book Contextual Guide",
    "Draft Checkpoint",
    "Meaning Before Application",
    "Grasping the Scene",
    "Application: Write the Report and Act",
    "The verdict",
    "The action",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_body(doc, text, italic=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.keep_together = False
    run = p.add_run("\u00a0" + text)
    set_run_font(run, size=11, color=NAVY, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=11, color=NAVY)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PAPER)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E74B5")
    p_bdr.append(left)
    p_pr.append(p_bdr)
    label_run = p.add_run(label.upper() + "  ")
    set_run_font(label_run, size=9, color=BLUE, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=11, color=NAVY, italic=True)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    label_style = doc.styles.add_style("Reference Label", WD_STYLE_TYPE.PARAGRAPH)
    label_style.font.name = "Calibri"
    label_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    label_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    label_style.font.size = Pt(9)
    label_style.font.bold = True
    label_style.font.color.rgb = BLUE
    label_style.paragraph_format.space_before = Pt(8)
    label_style.paragraph_format.space_after = Pt(3)
    label_style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = NAVY
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        p = header.paragraphs[0]
        p.text = "GENERAL HERMENEUTICS  |  INFOGRAPHIC SOURCE EDITION"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        add_page_field(footer.paragraphs[0])


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(112)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GENERAL HERMENEUTICS")
    set_run_font(run, size=11, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Complete Text for Infographic Creation")
    set_run_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Overview, Preparation, Observation, Interpretation, Imagination, and Application")
    set_run_font(run, size=13, color=MUTED)

    add_callout(
        doc,
        "Source edition",
        "Worked passage examples have been omitted so every section can serve as reusable infographic source material.",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INCLUDED SECTIONS")
    set_run_font(run, size=9, color=BLUE, bold=True)

    for item in SECTIONS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        set_run_font(run, size=11, color=NAVY)

    doc.add_page_break()


def is_all_caps_heading(line):
    letters = [ch for ch in line if ch.isalpha()]
    return bool(letters) and all(ch.isupper() for ch in letters) and len(line) <= 90


def add_document_content(doc, markdown):
    lines = markdown.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## "))
    lines = lines[start:]
    index = 0
    first_section = True
    list_mode = False

    while index < len(lines):
        line = lines[index].strip()

        if not line:
            list_mode = False
            index += 1
            continue

        if line.startswith("## "):
            if not first_section:
                doc.add_page_break()
            heading = line[3:].strip()
            p = doc.add_paragraph(heading, style="Heading 1")
            p.paragraph_format.page_break_before = False
            first_section = False
            index += 1
            continue

        if line == "General Hermeneutics: Overview":
            p = doc.add_paragraph(style="Reference Label")
            p.add_run("OVERVIEW")
            index += 1
            continue

        if line == "DETECTIVE LENS" and index + 1 < len(lines):
            add_callout(doc, line, lines[index + 1].strip())
            index += 2
            continue

        if line == "SECTION" and index + 10 < len(lines):
            context_header = [item.strip() for item in lines[index + 1 : index + 11]]
            if context_header == [
                "STEP 1",
                "BOOK",
                "STEP 2",
                "SAME-AUTHOR",
                "STEP 3",
                "CANONICAL",
                "STEP 4",
                "REDEMPTIVE-HISTORICAL",
                "STEP 5",
                "1",
            ]:
                index += 10
                continue

        if line.isdigit() and index + 1 < len(lines) and lines[index + 1].strip():
            title = lines[index + 1].strip()
            doc.add_paragraph(f"Step {line}: {title}", style="Heading 2")
            index += 2
            continue

        if line.startswith("STEP ") and line[5:].isdigit() and index + 1 < len(lines):
            title = lines[index + 1].strip()
            doc.add_paragraph(f"{line.title()}: {title}", style="Heading 2")
            index += 2
            continue

        normalized = line.upper()
        if normalized in MINOR_LABELS:
            p = doc.add_paragraph(style="Reference Label")
            p.add_run(line)
            list_mode = normalized in {"HOW TO BEGIN", "HOW TO DO IT", "HOW TO ENTER THE SCENE"}
            index += 1
            continue

        if normalized in MAJOR_LABELS or normalized.startswith(("A. ", "B. ", "C. ", "D. ")):
            doc.add_paragraph(line, style="Heading 2")
            list_mode = normalized in {"S.P.A.C.E.P.E.T.S.", "HOW TO DO A WORD STUDY"}
            index += 1
            continue

        if line in TITLE_HEADINGS:
            doc.add_paragraph(line, style="Heading 2")
            index += 1
            continue

        if line in {"Ask", "Look For", "Avoid", "The verdict", "The action"}:
            p = doc.add_paragraph(style="Reference Label")
            p.add_run(line)
            index += 1
            continue

        if line == "SELECT A BOOK":
            doc.add_paragraph("Books in the contextual guide", style="Heading 3")
            books = []
            cursor = index + 1
            while cursor < len(lines):
                candidate = lines[cursor].strip()
                if candidate == "Genesis" and books:
                    break
                books.append(candidate)
                cursor += 1
            add_body(doc, "; ".join(book for book in books if book))
            index = cursor
            continue

        if list_mode:
            add_bullet(doc, line)
        elif is_all_caps_heading(line):
            p = doc.add_paragraph(style="Reference Label")
            p.add_run(line)
        elif line.startswith("— "):
            add_body(doc, line, italic=True)
        else:
            add_body(doc, line)
        index += 1


def build():
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)

    props = doc.core_properties
    props.title = "General Hermeneutics Complete Text for Infographic Creation"
    props.subject = "Overview and five phases of General Hermeneutics"
    props.author = "My Bible Explorer"
    props.keywords = "hermeneutics, Bible study, infographic, interpretation"

    add_cover(doc)
    add_document_content(doc, markdown)
    configure_sections(doc)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
